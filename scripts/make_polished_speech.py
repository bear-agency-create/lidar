# -*- coding: utf-8 -*-
"""Lightly polished AirPorter speech — RU + EN on Desktop."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DESKTOP = r"c:\Users\user\Desktop"

RU = [
    ("Слайд 1. Знакомство с проектом", [
        "AirPorter — мобильный робот-ассистент для аэропорта",
        "Добрый день! Мы представляем проект AirPorter — мобильного робота-ассистента для аэропорта.",
        "Цель проекта — сделать пребывание пассажиров в аэропорту более комфортным. Робот помогает ориентироваться в терминале, сопровождает до нужной точки и может перевозить багаж, снижая физическую нагрузку и стресс — особенно для тех, кто оказался в аэропорту впервые.",
    ]),
    ("Слайд 2. Проблема", [
        "Современный аэропорт — сложная система с большим объёмом информации, сервисов и зон. Для многих пассажиров, особенно иностранных или редко летающих, ориентирование становится серьёзной проблемой.",
        "Мы выделили три основные сложности:",
        "• физическая нагрузка при перевозке багажа;",
        "• сложная навигация внутри терминала;",
        "• языковой барьер при получении информации.",
        "Именно эти проблемы легли в основу нашего проекта.",
    ]),
    ("Слайд 3. Решение", [
        "Мы предлагаем робота-ассистента, который помогает пассажиру пройти путь от стартовой точки до пункта назначения с учётом его потребностей.",
        "Работа системы состоит из нескольких этапов.",
        "Первый этап — сканирование билета. Робот получает основную информацию о пассажире и его рейсе.",
        "Второй этап — отображение информации. На экране показываются номер рейса, время посадки, выход и доступные сервисы.",
        "Далее возможны три сценария работы.",
        "Первый сценарий. Пассажиру нужна только информация. Робот показывает маршрут и необходимые данные, после чего сессия завершается.",
        "Второй сценарий. Пассажиру требуется сопровождение. Робот предлагает маршрут с учётом выбранных точек: регистрация, туалет, Duty Free, кафе или выход на посадку. После выбора маршрута робот сопровождает пассажира.",
        "Третий сценарий. Пассажиру требуется помощь только перевозкой багажа. Робот берёт вещи на себя и проходит все этапы проверки багажа; клиенту остаётся пройти личный досмотр и сесть в самолёт.",
        "Во втором и третьем сценариях предусмотрена работа с багажом. Если полное сопровождение не нужно, робот также может использоваться как временная точка хранения вещей.",
    ]),
    ("Слайд 4. Аппаратная платформа", [
        "Для реализации проекта мы собрали первый рабочий прототип.",
        "Основной вычислительный модуль — Raspberry Pi 5 (16 ГБ). На нём работают ROS 2 Jazzy, обработка данных лидара и камеры, а также высокоуровневая логика.",
        "Управление двигателями выполняет Arduino Mega 2560: она обрабатывает сигналы энкодеров и обеспечивает движение платформы.",
        "Используются драйверы L298N, моторы JGB37-520, колёса Mecanum, сенсорный экран 10,1\", камера и лидар COIN D6.",
    ]),
    ("Слайд 5. Навигация", [
        "Реализована система построения карты помещения с помощью лидара. Робот определяет своё положение, учитывает габариты корпуса и данные энкодеров, строит маршрут и объезжает препятствия. Есть режим проезда по заданным точкам и операторский контроль.",
    ]),
    ("Слайд 6. Интерфейс", [
        "На главном экране — основные сервисы аэропорта: регистрация, посадка, выдача багажа, информационная стойка и другие точки терминала.",
        "Поддерживаются русский, английский и татарский языки. Интерфейс выполнен в тематике аэропорта и рассчитан на простое использование пассажиром.",
    ]),
    ("Слайд 7. Безопасность", [
        "Реализованы автоматическая остановка при потере связи, контроль маршрута, распознавание препятствий, защитный корпус из оргстекла, небольшой дорожный просвет и световая индикация состояния робота.",
    ]),
    ("Слайд 8. Первый прототип", [
        "Сейчас мы представляем первый прототип команды: интерактивное меню на экране, проезд по точкам, ориентирование в пространстве и возможность перевозить груз.",
    ]),
    ("Слайд 9. Экспертная оценка", [
        "Проект был представлен экспертам авиационной отрасли и получил положительную оценку. Представители аэропорта выразили заинтересованность в дальнейшем тестировании разработки.",
    ]),
    ("Слайд 10. Экономическая часть", [
        "Стоимость текущего прототипа — около 90 000 рублей. Модульная конструкция позволяет модернизировать систему без полной замены оборудования.",
    ]),
    ("Слайд 11. Планы на будущее", [
        "Наши планы масштабные и амбициозные: впереди большая и интересная работа над проектом.",
        "Первое — усилить конструкцию металлическими профилями. Это повысит грузоподъёмность и устойчивость к ударам.",
        "Потребуются изменения и в колёсной базе: планируем промышленные металлические колёса Mecanum.",
        "Важно сохранить безопасность и аккуратный внешний вид: при переходе на металл добавим больше защитных и косметических элементов корпуса.",
        "По электронике планируем поэтапно переходить на промышленные компоненты, чтобы быстрее выйти на испытания в реальных условиях аэропорта.",
        "В навигации одна из ключевых задач — усилить сенсорику, в том числе промышленной камерой на 360° и дополнительными модулями.",
        "Также планируем удобные интерфейсы для пассажиров и для команды разработки.",
        "Для пользователей — приложение, где видно, где находится робот, на каком он этапе обслуживания, и при необходимости — уведомления о проблемах с багажом.",
        "Для разработчиков — панель флота: состояние систем, занятость, возможные неисправности.",
        "Так команда сможет точечно находить проблемы, устранять их удалённо или направлять робота в специализированный хаб для ремонта.",
        "Хаб станет станцией зарядки и местом хранения роботов, ожидающих обслуживания, откуда их можно забирать в мастерскую для ремонта и очистки корпуса.",
        "Отдельно планируем добавить подогрев и охлаждение напитков в уже имеющийся подстаканник.",
        "Заключительный этап — испытания системы в реальных сценариях аэропорта.",
    ]),
    ("Слайд 13. Заключение", [
        "Спасибо за внимание!",
        "Команда AirPorter",
    ]),
]

EN = [
    ("Slide 1. Introduction to the project", [
        "AirPorter — a mobile robot assistant for the airport",
        "Good afternoon! We present the AirPorter project — a mobile robot assistant for the airport.",
        "The goal of the project is to make passengers’ time at the airport more comfortable. The robot helps people find their way around the terminal, escorts them to the right location, and can carry luggage — reducing physical strain and stress, especially for first-time visitors.",
    ]),
    ("Slide 2. The problem", [
        "A modern airport is a complex system with a large amount of information, services, and zones. For many passengers — especially international travelers or those who fly rarely — finding one’s way is a serious challenge.",
        "We identified three main difficulties:",
        "• physical strain when carrying luggage;",
        "• complex navigation inside the terminal;",
        "• the language barrier when obtaining information.",
        "These problems formed the foundation of our project.",
    ]),
    ("Slide 3. The solution", [
        "We propose a robot assistant that helps a passenger travel from a starting point to a destination according to their needs.",
        "The system works in several stages.",
        "Stage one — ticket scanning. The robot receives key information about the passenger and their flight.",
        "Stage two — displaying information. The screen shows the flight number, boarding time, gate, and available services.",
        "Then three operating scenarios are possible.",
        "Scenario one. The passenger only needs information. The robot shows the route and required data, and the session ends.",
        "Scenario two. The passenger needs escorting. The robot suggests a route via selected points: check-in, restrooms, Duty Free, café, or the boarding gate. After the route is chosen, the robot escorts the passenger.",
        "Scenario three. The passenger needs help only with luggage transport. The robot takes the bags and goes through all baggage-screening stages; the passenger only needs to complete personal security screening and board the aircraft.",
        "In the second and third scenarios the robot can handle luggage. If full escorting is not required, it can also serve as a temporary place to store belongings.",
    ]),
    ("Slide 4. Hardware platform", [
        "To implement the project, we built the first working prototype.",
        "The main computing module is a Raspberry Pi 5 (16 GB). It runs ROS 2 Jazzy, processes lidar and camera data, and handles high-level logic.",
        "Motor control is performed by an Arduino Mega 2560, which processes encoder signals and drives the platform.",
        "We use L298N drivers, JGB37-520 motors, Mecanum wheels, a 10.1\" touchscreen, a camera, and a COIN D6 lidar.",
    ]),
    ("Slide 5. Navigation", [
        "We implemented indoor mapping with lidar. The robot estimates its pose, accounts for its footprint and encoder data, builds a route, and avoids obstacles. It supports waypoint travel and operator supervision.",
    ]),
    ("Slide 6. Interface", [
        "The home screen presents the main airport services: check-in, boarding, baggage claim, the information desk, and other terminal points.",
        "Russian, English, and Tatar are supported. The interface follows an airport visual theme and is designed to be simple for passengers.",
    ]),
    ("Slide 7. Safety", [
        "We implemented automatic stop on loss of connection, route monitoring, obstacle recognition, a protective acrylic body, low ground clearance, and status lighting.",
    ]),
    ("Slide 8. First prototype", [
        "Today we present our team’s first prototype: an interactive on-screen menu, waypoint travel, spatial orientation, and the ability to carry a load.",
    ]),
    ("Slide 9. Expert feedback", [
        "The project was presented to aviation-industry experts and received positive feedback. Airport representatives expressed interest in further testing.",
    ]),
    ("Slide 10. Economics", [
        "The current prototype costs about 90,000 rubles. The modular design allows upgrades without fully replacing the hardware.",
    ]),
    ("Slide 11. Future plans", [
        "Our plans are large and ambitious: substantial and exciting work on the project still lies ahead.",
        "First, we will reinforce the structure with metal profiles to increase payload capacity and impact resistance.",
        "The wheelbase will also change: we plan industrial metal Mecanum wheels.",
        "Safety and appearance remain important — with the move to metal we will add more protective and cosmetic body details.",
        "On the electronics side, we plan a gradual shift to industrial-grade components to reach real airport trials sooner.",
        "In navigation, a key goal is stronger sensing, including a 360° industrial camera and additional modules.",
        "We will also develop convenient interfaces for passengers and for the development team.",
        "For users — an app showing where the robot is, what stage of service it is in, and baggage alerts when needed.",
        "For developers — a fleet dashboard with system health, occupancy, and possible faults.",
        "This will let the team pinpoint issues, fix them remotely, or send a robot to a specialized hub for repair.",
        "The hub will be a charging station and holding area for robots awaiting service, from which they can be taken to the workshop for repair and cleaning.",
        "We also plan heating and cooling for drinks in the existing cup holder.",
        "The final stage is testing the system in real airport scenarios.",
    ]),
    ("Slide 13. Closing", [
        "Thank you for your attention!",
        "The AirPorter Team",
    ]),
]


def write_docx(path: str, title: str, sections):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for slide_title, paras in sections:
        doc.add_heading(slide_title, level=1)
        for text in paras:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

    doc.save(path)
    print("saved", path)


write_docx(
    DESKTOP + r"\AirPorter Выступление (RU, правленый).docx",
    "AirPorter — текст выступления",
    RU,
)
write_docx(
    DESKTOP + r"\AirPorter Speech (EN, polished).docx",
    "AirPorter — presentation speech",
    EN,
)
print("done")
