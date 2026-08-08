# -*- coding: utf-8 -*-
"""Combine the polished RU and EN speeches into one final Word document."""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

DESKTOP = Path.home() / "Desktop"
TEMPLATE = DESKTOP / "AirPorter Выступление(рус.).docx"
RU = DESKTOP / "AirPorter Выступление (RU, правленый).docx"
EN = DESKTOP / "AirPorter Speech (EN, polished).docx"
OUTPUT = DESKTOP / "AirPorter — Final Speech (RU + EN).docx"


def clear_body(document: Document) -> None:
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def append_document_paragraphs(destination: Document, source: Document) -> None:
    body = destination._element.body
    section_properties = body.sectPr
    for paragraph in source.paragraphs:
        section_properties.addprevious(deepcopy(paragraph._p))


def main() -> None:
    for path in (TEMPLATE, RU, EN):
        if not path.is_file():
            raise FileNotFoundError(path)

    result = Document(TEMPLATE)
    clear_body(result)
    russian = Document(RU)
    english = Document(EN)

    append_document_paragraphs(result, russian)
    separator = result.add_paragraph(style="Normal")
    separator.add_run().add_break(WD_BREAK.PAGE)
    append_document_paragraphs(result, english)

    result.core_properties.title = "AirPorter — Final Speech (Russian and English)"
    result.core_properties.subject = "Polished presentation speech"
    result.core_properties.author = "AirPorter Team"

    expected = [
        paragraph.text
        for document in (russian, english)
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    actual = [paragraph.text for paragraph in result.paragraphs if paragraph.text.strip()]
    if actual != expected:
        raise RuntimeError("Combined document text does not match polished sources")

    temporary = OUTPUT.with_name(f"{OUTPUT.stem}.tmp.docx")
    result.save(temporary)
    temporary.replace(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
