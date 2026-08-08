# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)

sections = [
    (
        "Slide 1. Introduction to the project",
        """AirPorter — a mobile robot assistant for the airport

Good afternoon! We present the AirPorter project — a mobile robot assistant for the airport.

The goal of the project is to make passengers’ time at the airport more comfortable. The robot helps people find their way around the terminal, escorts them to the right location, and can also carry luggage, reducing physical strain and stress — especially for those visiting an airport for the first time.""",
    ),
    (
        "Slide 2. The problem",
        """A modern airport is a complex system with a large amount of information, services, and zones. For most passengers — especially international travelers or those who fly rarely — finding one’s way becomes a serious challenge.

We identified three main difficulties:

• physical strain when carrying luggage;

• complex navigation inside the terminal;

• the language barrier when obtaining information.

These problems formed the foundation of our project.""",
    ),
    (
        "Slide 3. The solution",
        """We propose using an autonomous robot that escorts a passenger from a starting point to a destination while taking their needs into account.

The system works in several stages.

Stage one — ticket scanning. The robot receives key information about the passenger and their flight.

Stage two — displaying information. The screen shows the flight number, boarding time, gate, and available services.

Then three operating scenarios are possible.

Scenario one. The passenger only needs information. The robot shows the route and required data, and the session ends.

Scenario two. The passenger needs escorting. The robot suggests an optimal route via selected points: check-in, restrooms, Duty Free, café, or the boarding gate. After the route is chosen, the robot escorts the passenger.

Scenario three. The passenger needs help with luggage. The robot takes the bags and escorts the passenger through all stages up to boarding.

In all scenarios except the first, luggage transport is available. If escorting is not required, the robot can also serve as a temporary place to store belongings.""",
    ),
    (
        "Slide 4. Hardware platform",
        """To implement the project, we built the first working prototype.

The main computing module is a Raspberry Pi 5 (16 GB). It runs ROS 2 Jazzy, processes lidar and camera data, and handles high-level logic.

Motor control is performed by an Arduino Mega 2560, which processes encoder signals and provides precise platform motion.

We use L298N drivers, JGB37-520 motors, Mecanum wheels, a 10.1" touchscreen, a camera, and a COIN D6 lidar.""",
    ),
    (
        "Slide 5. Navigation",
        """We implemented indoor mapping using lidar. The robot estimates its own pose, accounts for its footprint and encoder data, builds an optimal path, and avoids obstacles.""",
    ),
    (
        "Slide 6. Interface",
        """The home screen presents the main airport services: check-in, boarding, baggage claim, the information desk, and other terminal points.

Russian, English, and Tatar are supported. The interface follows an airport visual theme and is designed to be as simple as possible for passengers.""",
    ),
    (
        "Slide 7. Safety",
        """We implemented automatic stop on loss of connection, route monitoring, obstacle recognition, a protective acrylic body, minimal ground clearance, and status lighting.""",
    ),
    (
        "Slide 8. First prototype",
        """Today we present our team’s first prototype, which already includes an interactive on-screen menu, waypoint travel, spatial orientation, and load-bearing capability.""",
    ),
    (
        "Slide 9. Expert feedback",
        """The project was presented to aviation-industry experts and received positive feedback. Airport representatives expressed interest in further testing of the development.""",
    ),
    (
        "Slide 10. Economics",
        """The cost of the current prototype is about 90,000 rubles. The modular design allows the system to be upgraded without fully replacing the hardware.""",
    ),
    (
        "Slide 11. Future plans",
        """Our future plans are large and ambitious: a great deal of interesting work still lies ahead on this project.

First, we will reinforce the structure with metal profiles. This will help the robot carry heavier loads in the body and better withstand impacts and drops.

The body needs changes not only in structure but also in the wheels. We plan to install industrial Mecanum metal wheels.

Appearance and safety also matter. Because we will use metal profiles, the team will add more protective cosmetic details.

As for electronics, we intend to replace components with industrial-grade parts so the robot can be tested sooner in real airport conditions.

Navigation is one of the key challenges. We plan to add a 360-degree industrial camera and more.

We will also develop a convenient interface for using and tracking the robot — both for passengers and for developers.

Users will get an app showing where the robot is, what stage it is at, and, when needed, alerts about baggage issues.

On the developer side there will be a fleet dashboard: system health, occupancy, possible problems, and so on.

This will let the team pinpoint issues on specific robots, fix them remotely, or call them into a specialized hub for physical repair.

That hub will be a charging station and a holding area for robots awaiting repair, from which the team can take them to the workshop for maintenance and cleaning when needed.

We will also add heating and cooling for drinks in the existing cup holder.

The final stage of the project is testing the system in real airport situations.""",
    ),
    (
        "Slide 13. Closing",
        """Thank you for your attention!

The AirPorter Team""",
    ),
]

for title, body in sections:
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(6)

out = r"c:\Users\user\Desktop\AirPorter Speech (EN).docx"
doc.save(out)
print("saved", out)
