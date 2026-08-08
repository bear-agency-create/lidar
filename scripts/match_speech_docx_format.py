# -*- coding: utf-8 -*-
"""Apply the original Russian speech document's formatting to polished files."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DESKTOP = Path.home() / "Desktop"
SOURCE = DESKTOP / "AirPorter Выступление(рус.).docx"
TARGETS = [
    DESKTOP / "AirPorter Speech (EN, polished).docx",
    DESKTOP / "AirPorter Выступление (RU, правленый).docx",
]


def slide_number(text: str) -> int | None:
    match = re.match(r"^(?:Slide|Слайд)\s+(\d+)", text.strip(), re.IGNORECASE)
    return int(match.group(1)) if match else None


def nonblank_with_gaps(doc: Document) -> list[tuple[str, int]]:
    result: list[tuple[str, int]] = []
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip():
            continue
        gap = 0
        cursor = index + 1
        while cursor < len(paragraphs) and not paragraphs[cursor].text.strip():
            gap += 1
            cursor += 1
        result.append((paragraph.text, gap))
    return result


def sections(items: list[tuple[str, int]]) -> tuple[list[tuple[str, int]], dict[int, list[tuple[str, int]]]]:
    prefix: list[tuple[str, int]] = []
    grouped: dict[int, list[tuple[str, int]]] = {}
    current: int | None = None
    for item in items:
        number = slide_number(item[0])
        if number is not None:
            current = number
            grouped[current] = [item]
        elif current is None:
            prefix.append(item)
        else:
            grouped[current].append(item)
    return prefix, grouped


def map_content_and_gaps(source: Document, target: Document) -> list[tuple[str, int]]:
    source_prefix, source_sections = sections(nonblank_with_gaps(source))
    target_prefix, target_sections = sections(nonblank_with_gaps(target))
    output: list[tuple[str, int]] = []

    # The polished documents have one cover line that the original did not have.
    for text, _ in target_prefix:
        output.append((text, 1))

    for number, target_items in target_sections.items():
        source_items = source_sections.get(number, [])
        if len(source_items) == len(target_items):
            output.extend((target_item[0], source_item[1]) for target_item, source_item in zip(target_items, source_items))
        elif source_items:
            # Slide 10 was polished from two body paragraphs into one.
            for index, target_item in enumerate(target_items):
                if index == 0:
                    gap = source_items[0][1]
                elif index == len(target_items) - 1:
                    gap = source_items[-1][1]
                else:
                    mapped = round(index * (len(source_items) - 1) / max(1, len(target_items) - 1))
                    gap = source_items[mapped][1]
                output.append((target_item[0], gap))
        else:
            output.extend((text, gap) for text, gap in target_items)
    return output


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def add_formatted_paragraph(doc: Document, text: str, language: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    run.bold = False
    run.italic = False
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")
    lang = run._element.get_or_add_rPr().find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(lang)
    lang.set(qn("w:val"), language)


def format_target(path: Path) -> None:
    source = Document(SOURCE)
    target = Document(path)
    expected_text = [paragraph.text for paragraph in target.paragraphs if paragraph.text.strip()]
    content = map_content_and_gaps(source, target)

    result = Document(SOURCE)
    clear_document_body(result)
    language = "en-US" if path.name.startswith("AirPorter Speech") else "ru-RU"
    for text, gap_after in content:
        add_formatted_paragraph(result, text, language)
        for _ in range(gap_after):
            result.add_paragraph(style="Normal")

    actual_text = [paragraph.text for paragraph in result.paragraphs if paragraph.text.strip()]
    if actual_text != expected_text:
        raise RuntimeError(f"Text changed while formatting {path.name}")

    temporary = path.with_name(f"{path.stem}.formatting.tmp.docx")
    result.save(temporary)
    temporary.replace(path)
    print(f"formatted: {path}")


def main() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    for target in TARGETS:
        if not target.is_file():
            raise FileNotFoundError(target)
        format_target(target)


if __name__ == "__main__":
    main()
